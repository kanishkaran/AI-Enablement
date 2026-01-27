"""
Enhanced MCP Server for Google Drive Documents
Supports both Google Docs and uploaded DOCX files
"""

from fastmcp import FastMCP
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import os.path
import pickle
import io
from docx import Document

# Initialize FastMCP server
mcp = FastMCP("Google Drive Documents Assistant")

# Google API Scopes - read-only access
SCOPES = [
    'https://www.googleapis.com/auth/documents.readonly',
    'https://www.googleapis.com/auth/drive.readonly'
]

def get_google_services():
    """Authenticate and return Google Docs and Drive services"""
    creds = None
    
    # Check for existing token
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    
    # Refresh or create new credentials
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'Week-6/guardrails/Agent-Task/src/tools/mcp_tool/credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Save credentials for future use
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    
    docs_service = build('docs', 'v1', credentials=creds)
    drive_service = build('drive', 'v3', credentials=creds)
    
    return docs_service, drive_service

def extract_google_doc_text(document):
    """Extract all text content from a Google Doc"""
    text_parts = []
    
    content = document.get('body', {}).get('content', [])
    
    for element in content:
        # Handle paragraphs
        if 'paragraph' in element:
            for text_element in element['paragraph'].get('elements', []):
                if 'textRun' in text_element:
                    text_parts.append(text_element['textRun'].get('content', ''))
        
        # Handle tables
        elif 'table' in element:
            for row in element['table'].get('tableRows', []):
                for cell in row.get('tableCells', []):
                    for cell_content in cell.get('content', []):
                        if 'paragraph' in cell_content:
                            for text_element in cell_content['paragraph'].get('elements', []):
                                if 'textRun' in text_element:
                                    text_parts.append(text_element['textRun'].get('content', ''))
    
    return ''.join(text_parts)

def extract_docx_text(drive_service, file_id):
    """Extract text from a DOCX file in Google Drive"""
    try:
        request = drive_service.files().get_media(fileId=file_id)
        file_buffer = io.BytesIO()
        downloader = MediaIoBaseDownload(file_buffer, request)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_buffer.seek(0)
        doc = Document(file_buffer)
        
        # Extract all paragraphs
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    except Exception as e:
        return f"[Error extracting DOCX content: {str(e)}]"

def get_document_content_by_type(docs_service, drive_service, file_id, mime_type):
    """Get document content based on file type"""
    if mime_type == 'application/vnd.google-apps.document':
        # Native Google Doc
        document = docs_service.documents().get(documentId=file_id).execute()
        return extract_google_doc_text(document)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # DOCX file
        return extract_docx_text(drive_service, file_id)
    else:
        return "[Unsupported document type]"


@mcp.tool()
def list_recent_docs(count: int = 15) -> str:
    """
    List recently modified documents (both Google Docs and DOCX files).
    Useful for browsing available documents.
    
    Args:
        count: Number of recent documents to list (default: 15)
    
    Returns:
        List of recent documents with titles, types, and IDs
    """
    try:
        _, drive_service = get_google_services()
        
        # Query for both Google Docs and DOCX files
        query = (
            "mimeType='application/vnd.google-apps.document' or "
            "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"
        )
        
        results = drive_service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name, mimeType, modifiedTime)',
            pageSize=count,
            orderBy='modifiedTime desc'
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            return "No documents found in your Drive"
        
        response_parts = [f"Recent {len(files)} document(s):\n"]
        
        for idx, file in enumerate(files, 1):
            file_type = "Google Doc" if file['mimeType'] == 'application/vnd.google-apps.document' else "DOCX"
            
            response_parts.append(
                f"\n{idx}. {file['name']} [{file_type}]\n"
                f"   Document ID: {file['id']}\n"
                f"   Last Modified: {file.get('modifiedTime', 'Unknown')}\n"
            )
        
        return ''.join(response_parts)
    
    except Exception as e:
        return f"Error listing documents: {str(e)}"

@mcp.tool()
def search_and_retrieve(query: str) -> str:
    """
    Combined search and retrieval - finds the most relevant document related to health insurance and term life insurance
    and returns its full content. Works with both Google Docs and DOCX files.
    
    Args:
        query: Search term to find the document
    
    Returns:
        Full content of the most relevant document
    """
    try:
        docs_service, drive_service = get_google_services()
        
        # Search for documents
        drive_query = (
            f"(mimeType='application/vnd.google-apps.document' or "
            f"mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') "
            f"and fullText contains '{query}'"
        )
        
        results = drive_service.files().list(
            q=drive_query,
            spaces='drive',
            fields='files(id, name, mimeType)',
            pageSize=1,
            orderBy='modifiedTime desc'
        ).execute()
        
        files = results.get('files', [])
        
        if not files:
            return f"No documents found matching: '{query}'"
        
        # Get the most recent matching document
        doc_id = files[0]['id']
        doc_name = files[0]['name']
        mime_type = files[0]['mimeType']
        
        # Retrieve full content
        content = get_document_content_by_type(
            docs_service, drive_service, doc_id, mime_type
        )
        
        file_type = "Google Doc" if mime_type == 'application/vnd.google-apps.document' else "DOCX"
        
        return f"=== {doc_name} [{file_type}] ===\nDocument ID: {doc_id}\n\n{content}"
    
    except Exception as e:
        return f"Error in search and retrieval: {str(e)}"

if __name__ == "__main__":
    
    mcp.run()